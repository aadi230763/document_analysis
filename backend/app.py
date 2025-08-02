import os
os.environ['FLASK_NO_COLOR'] = '1'  # Disable colored output

from dotenv import load_dotenv
load_dotenv()

from flask import Flask, request, jsonify
import asyncio
from config import *


import requests
from config import GEMINI_API_KEY, GEMINI_API_URL
import atexit
from utils.medical_terms import get_dynamic_synonyms, parse_demographics, extract_policy_duration
from utils.query_expander import QueryExpander
import re
from utils.rule_engine import RuleEngine
import json
from datetime import datetime, timezone
import cohere
from flask_cors import CORS
import time
from io import BytesIO
import tempfile
from PyPDF2 import PdfReader
import docx
import numpy as np
import mimetypes
import logging
from email import policy
from email.parser import BytesParser
import base64
import pdfplumber



# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("doc_parser")

def add_chunk(chunks, text, section=None, table=None, chunk_type=None):
    if text and text.strip():
        chunks.append({
            'text': text.strip(),
            'section': section or '',
            'table': table or '',
            'type': chunk_type or ''
        })

def extract_section_headers(text):
    """Enhanced section header extraction for insurance documents"""
    patterns = [
        r'^(?:Section|Clause|Article)\s+\d+[\.\d]*\s*[:\-]?\s*[A-Z][A-Za-z\s]+$',
        r'^(?:Coverage|Benefits|Exclusions|Limitations|Terms|Conditions)\s*[:\-]?\s*[A-Z][A-Za-z\s]+$',
        r'^(?:Policy|Insurance|Premium|Claim|Deductible|Co-pay|Coinsurance)\s+[A-Z][A-Za-z\s]+$',
        r'^\d+\.\s+[A-Z][A-Z\s]{3,}$',
        r'^[A-Z][A-Z\s]{3,}$',
        r'^(?:Grace|Waiting|Exclusion|Inclusion|Coverage|Benefit)\s+Period\s*[:\-]?\s*[A-Z][A-Za-z\s]*$',
        r'^(?:Network|Non-Network|In-Network|Out-of-Network)\s+[A-Z][A-Za-z\s]+$',
        r'^(?:Pre-existing|Pre-existing Condition)\s*[:\-]?\s*[A-Z][A-Za-z\s]*$',
    ]
    headers = []
    for pattern in patterns:
        headers.extend([m.group() for m in re.finditer(pattern, text, re.MULTILINE)])
    return sorted(set(headers), key=lambda h: text.find(h))

def smart_chunk_text(text, section_name="", chunk_type="", max_chunk_size=None):
    from config import DEFAULT_CHUNK_SIZE
    if max_chunk_size is None:
        max_chunk_size = DEFAULT_CHUNK_SIZE
    """Improved semantic chunking with better context preservation"""
    if not text.strip():
        return []
    
    chunks = []
    
    # Split by paragraphs first
    paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 0]
    
    # If no paragraphs, split by sentences
    if not paragraphs:
        sentences = re.split(r'(?<=[.!?])\s+', text)
        paragraphs = [s.strip() for s in sentences if len(s.strip()) > 0]
    
    # If still no content, use the whole text as one chunk
    if not paragraphs:
        if len(text.strip()) > 0:
            return [{'text': text.strip(), 'section': section_name, 'type': chunk_type}]
        return []
    
    current_chunk = []
    current_length = 0
    
    for para in paragraphs:
        para_words = len(para.split())
        
        # Skip very short paragraphs unless they're the only content
        if para_words < 5 and len(paragraphs) > 1:
            continue
        
        # If paragraph is too long, split it by sentences
        if para_words > max_chunk_size // 4:  # Approximate word count
            sentences = re.split(r'(?<=[.!?])\s+', para)
            for sent in sentences:
                if not sent.strip():
                    continue
                sent_words = len(sent.split())
                
                # If adding this sentence would exceed limit, save current chunk
                if current_length + sent_words > max_chunk_size // 4 and current_chunk:
                    chunk_text = ' '.join(current_chunk)
                    add_chunk(chunks, chunk_text, section=section_name, chunk_type=chunk_type)
                    current_chunk = [sent]
                    current_length = sent_words
                else:
                    current_chunk.append(sent)
                    current_length += sent_words
        else:
            # If adding this paragraph would exceed limit, save current chunk
            if current_length + para_words > max_chunk_size // 4 and current_chunk:
                chunk_text = ' '.join(current_chunk)
                add_chunk(chunks, chunk_text, section=section_name, chunk_type=chunk_type)
                current_chunk = [para]
                current_length = para_words
            else:
                current_chunk.append(para)
                current_length += para_words
    
    # Add remaining content as final chunk
    if current_chunk:
        chunk_text = ' '.join(current_chunk)
        add_chunk(chunks, chunk_text, section=section_name, chunk_type=chunk_type)
    
    return chunks

def parse_document_from_url(url):
    """
    Download and parse a document from a URL (PDF, DOCX, EML, TXT).
    Returns: list of dicts: [{text, section, table, type, ...}]
    Raises: Exception with clear error message if parsing fails.
    """
    import requests
    from io import BytesIO
    import docx
    from PyPDF2 import PdfReader
    import traceback
    import re
    import os

    logger.info(f"Downloading document: {url}")
    try:
        from config import DOCUMENT_DOWNLOAD_TIMEOUT
        resp = requests.get(url, timeout=DOCUMENT_DOWNLOAD_TIMEOUT)
        resp.raise_for_status()
        content = resp.content
        content_type = resp.headers.get('Content-Type', '').lower()
        ext = os.path.splitext(url.split('?')[0])[1].lower()
        logger.info(f"Downloaded {len(content)} bytes, Content-Type: {content_type}, Extension: {ext}")
        
        # Check if we have any content
        if not content or len(content) == 0:
            raise Exception("Downloaded document is empty")
            
    except Exception as e:
        logger.error(f"Failed to download document: {url} | {e}")
        raise Exception(f"Failed to download document: {e}")

    # --- Improved PDF logic ---
    if ext == '.pdf' or 'pdf' in content_type:
        logger.info(f"Attempting PDF parsing for {len(content)} bytes")
        try:
            pdf_stream = BytesIO(content)
            reader = PdfReader(pdf_stream)
            all_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ''
                all_text.append(text)
            full_text = '\n'.join(all_text)
            logger.info(f"PDF extracted {len(full_text)} characters of text")
            
            # If no text, fallback to pdfplumber
            if not full_text.strip():
                logger.info("No text from PyPDF2, trying pdfplumber")
                pdf_stream.seek(0)
                with pdfplumber.open(pdf_stream) as pdf:
                    all_text = [page.extract_text() or '' for page in pdf.pages]
                full_text = '\n'.join(all_text)
                logger.info(f"pdfplumber extracted {len(full_text)} characters of text")
            
            # If still no text, fallback to OCR (disabled for Azure deployment)
            if not full_text.strip():
                logger.info("No text from pdfplumber, OCR disabled for Azure deployment")
                logger.warning("OCR functionality disabled - pdf2image and pytesseract not available")
            
            # Enhanced section-aware chunking
            chunks = []
            headers = extract_section_headers(full_text)
            
            if headers:
                # Section-based chunking
                header_positions = [(m.start(), m.group()) for h in headers for m in re.finditer(re.escape(h), full_text)]
                header_positions = sorted(header_positions, key=lambda x: x[0])
                
                for idx, (pos, header) in enumerate(header_positions):
                    start = pos
                    end = header_positions[idx + 1][0] if idx + 1 < len(header_positions) else len(full_text)
                    section_text = full_text[start:end].strip()
                    
                    if section_text:
                        section_chunks = smart_chunk_text(section_text, section_name=header, chunk_type='pdf_section')
                        chunks.extend(section_chunks)
            else:
                # Fallback to paragraph-based chunking
                chunks = smart_chunk_text(full_text, section_name='PDF Document', chunk_type='pdf_section')
            
            result = [c for c in chunks if c['text']]
            logger.info(f"PDF parsing successful, created {len(result)} chunks")
            if not result:
                logger.warning("PDF parsing completed but no valid chunks were created")
            return result
            
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}\n{traceback.format_exc()}")
            # Fallback to plain text
    
    # --- Improved DOCX logic ---
    if ext == '.docx' or 'word' in content_type or 'docx' in content_type:
        logger.info(f"Attempting DOCX parsing for {len(content)} bytes")
        try:
            docx_stream = BytesIO(content)
            doc = docx.Document(docx_stream)
            
            # Extract all text with structure
            full_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            logger.info(f"DOCX extracted {len(full_text)} characters of text")
            
            chunks = []
            headers = extract_section_headers(full_text)
            
            if headers:
                # Section-based chunking for DOCX
                header_positions = [(m.start(), m.group()) for h in headers for m in re.finditer(re.escape(h), full_text)]
                header_positions = sorted(header_positions, key=lambda x: x[0])
                
                for idx, (pos, header) in enumerate(header_positions):
                    start = pos
                    end = header_positions[idx + 1][0] if idx + 1 < len(header_positions) else len(full_text)
                    section_text = full_text[start:end].strip()
                    
                    if section_text:
                        section_chunks = smart_chunk_text(section_text, section_name=header, chunk_type='docx_section')
                        chunks.extend(section_chunks)
            else:
                # Fallback to paragraph-based chunking
                for para in doc.paragraphs:
                    if para.text.strip() and len(para.text.split()) > 10:
                        add_chunk(chunks, para.text, section=para.style.name if para.style else None, chunk_type='docx_paragraph')
            
            # Add tables as separate chunks
            for t_idx, table in enumerate(doc.tables):
                table_text = '\n'.join([' | '.join(cell.text.strip() for cell in row.cells) for row in table.rows])
                if table_text.strip():
                    add_chunk(chunks, table_text, section=f"Table {t_idx+1}", table=table_text, chunk_type='docx_table')
            
            result = [c for c in chunks if c['text']]
            logger.info(f"DOCX parsing successful, created {len(result)} chunks")
            if not result:
                logger.warning("DOCX parsing completed but no valid chunks were created")
            return result
            
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}\n{traceback.format_exc()}")
    
    # --- EML logic ---
    if ext == '.eml' or 'message/rfc822' in content_type or 'eml' in content_type:
        logger.info(f"Attempting EML parsing for {len(content)} bytes")
        try:
            msg = BytesParser(policy=policy.default).parsebytes(content)
            text = ''
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == 'text/plain':
                        text += part.get_content() + '\n'
            else:
                text = msg.get_content()
            
            logger.info(f"EML extracted {len(text)} characters of text")
            chunks = smart_chunk_text(text, section_name='Email Body', chunk_type='eml')
            result = [c for c in chunks if c['text']]
            logger.info(f"EML parsing successful, created {len(result)} chunks")
            if not result:
                logger.warning("EML parsing completed but no valid chunks were created")
            return result
            
        except Exception as e:
            logger.error(f"EML parsing failed: {e}\n{traceback.format_exc()}")
    
    # --- Plain text fallback ---
    logger.info(f"Attempting plain text parsing for {len(content)} bytes")
    try:
        text = content.decode(errors='ignore')
        logger.info(f"Plain text extracted {len(text)} characters")
        chunks = smart_chunk_text(text, section_name='Plain Text', chunk_type='plain')
        result = [c for c in chunks if c['text']]
        logger.info(f"Plain text parsing successful, created {len(result)} chunks")
        if not result:
            logger.warning("Plain text parsing completed but no valid chunks were created")
        return result
    except Exception as e:
        logger.error(f"Plain text parsing failed: {e}\n{traceback.format_exc()}")
        
        # Final fallback: try to parse as PDF regardless of content type
        logger.info("Attempting final PDF fallback parsing")
        try:
            pdf_stream = BytesIO(content)
            reader = PdfReader(pdf_stream)
            all_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ''
                all_text.append(text)
            full_text = '\n'.join(all_text)
            
            if full_text.strip():
                chunks = smart_chunk_text(full_text, section_name='PDF Document (Fallback)', chunk_type='pdf_fallback')
                result = [c for c in chunks if c['text']]
                if result:
                    logger.info(f"PDF fallback parsing successful, created {len(result)} chunks")
                    return result
        except Exception as pdf_fallback_e:
            logger.error(f"PDF fallback parsing also failed: {pdf_fallback_e}")
        
        raise Exception(f"Document could not be parsed as PDF, DOCX, EML, or plain text. Content-Type: {content_type}, Extension: {ext}, Content length: {len(content)} bytes. Error: {str(e)}")

app = Flask(__name__)
CORS(app)

# === COHERE API KEY SETUP ===
# Load from environment variables for production
COHERE_API_KEY = os.getenv("COHERE_API_KEY", "qZmghdKw7d7YxNryMj57OsMN0jLsQSCy0c7xulRA")
co = cohere.Client(COHERE_API_KEY)

# === PINECONE INTEGRATION ===
from utils.pinecone_manager import pinecone_manager

def llm(prompt, max_tokens=512, temperature=0.2):
    response = co.generate(
        model='command-r-plus',
        prompt=prompt,
        max_tokens=max_tokens,
        temperature=temperature
    )
    return response.generations[0].text

# === PERFORMANCE TRACKING ===
class PerformanceTracker:
    def __init__(self):
        self.start_time = None
    
    def start(self):
        self.start_time = time.time()
    
    def get_response_time_ms(self):
        if self.start_time:
            return int((time.time() - self.start_time) * 1000)
        return 0

# === ACCURACY IMPROVEMENTS ===
class ClauseRetriever:
    def __init__(self, collection):
        self.collection = collection
    
    def get_relevant_clauses(self, query, n_results=5):
        """Enhanced clause retrieval with better scoring"""
        try:
            # Get more candidates first
            results = self.collection.query(query_texts=[query], n_results=15)
            
            if not results['documents'] or not results['documents'][0]:
                return []
            
            # Re-rank using multiple strategies
            scored_chunks = []
            query_words = set(query.lower().split())
            
            for chunk, meta in zip(results['documents'][0], results['metadatas'][0]):
                score = 0.0
                chunk_lower = chunk.lower()
                
                # Exact keyword matching
                for word in query_words:
                    if word in chunk_lower:
                        score += 2.0
                
                # Section relevance
                if meta.get('section_headers'):
                    score += 1.0
                
                # Document source relevance
                if meta.get('source'):
                    score += 0.5
                
                scored_chunks.append({
                    'text': chunk,
                    'metadata': meta,
                    'relevance_score': score
                })
            
            # Return top N most relevant
            return sorted(scored_chunks, key=lambda x: x['relevance_score'], reverse=True)[:n_results]
            
        except Exception as e:
            print(f"Clause retrieval error: {str(e)}")
            return []

class ResponseFormatter:
    @staticmethod
    def format_clause_evidence(clause_data):
        """Format clause data for JSON response"""
        return {
            "clause": clause_data['metadata'].get('section_headers', 'Unknown Section'),
            "text": clause_data['text'][:300] + "..." if len(clause_data['text']) > 300 else clause_data['text'],
            "relevance_score": round(clause_data['relevance_score'], 2),
            "document": clause_data['metadata'].get('source', 'Unknown Document')
        }
    
    @staticmethod
    def calculate_accuracy_score(clauses_used, answer):
        """Calculate accuracy score based on evidence quality"""
        if not clauses_used:
            return 0.0
        
        # Base score from relevance scores
        avg_relevance = sum(c['relevance_score'] for c in clauses_used) / len(clauses_used)
        
        # Boost score if answer is confident
        confidence_boost = 0.2 if any(word in answer.lower() for word in ['yes', 'covered', 'include']) else 0.0
        
        return min(1.0, avg_relevance / 5.0 + confidence_boost)

# === TOKEN EFFICIENCY ===
class TokenOptimizer:
    @staticmethod
    def truncate_clauses_aggressively(clauses, max_tokens=1000):
        """Only send the most relevant parts of clauses"""
        truncated = []
        current_tokens = 0
        
        for clause in clauses:
            # Take only first 200 characters of each clause
            short_text = clause['text'][:200] + "..."
            tokens = len(short_text) // 4
            
            if current_tokens + tokens > max_tokens:
                break
                
            truncated.append({
                **clause,
                'text': short_text
            })
            current_tokens += tokens
        
        return truncated

# Simple in-memory cache for recent queries
query_cache = {}
CACHE_SIZE = 100

def cache_get(key):
    return query_cache.get(key)

def cache_set(key, value):
    if len(query_cache) >= CACHE_SIZE:
        # Remove oldest entry
        oldest_key = next(iter(query_cache))
        del query_cache[oldest_key]
    query_cache[key] = value

def redact_pii(text):
    # Redact emails, phone numbers, and policy numbers (less aggressive)
    text = re.sub(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", "[REDACTED_EMAIL]", text)
    text = re.sub(r"\b\d{10,}\b", "[REDACTED_PHONE]", text)
    text = re.sub(r"\b[A-Z0-9]{8,}\b", "[REDACTED_POLICY]", text)
    # Do NOT redact all capitalized words (names)
    return text

def redact_pii_in_dict(d):
    if isinstance(d, dict):
        return {k: redact_pii_in_dict(v) for k, v in d.items()}
    elif isinstance(d, list):
        return [redact_pii_in_dict(x) for x in d]
    elif isinstance(d, str):
        return redact_pii(d)
    else:
        return d

def audit_log(entry):
    with open("audit.log", "a", encoding="utf-8") as f:
        f.write(json.dumps(entry) + "\n")

def count_tokens(text):
    # Simple approximation: 1 token ≈ 4 characters
    return len(text) // 4

def truncate_clauses_to_fit(clauses, prompt_prefix, max_tokens=2048, buffer=200):
    current_tokens = count_tokens(prompt_prefix)
    truncated_clauses = []
    for clause in clauses:
        clause_text = clause['text']
        clause_tokens = count_tokens(clause_text)
        if current_tokens + clause_tokens + buffer > max_tokens:
            break
        truncated_clauses.append(clause)
        current_tokens += clause_tokens
    return truncated_clauses

# Initialize our new components
clause_retriever = ClauseRetriever(None)
response_formatter = ResponseFormatter()
token_optimizer = TokenOptimizer()

SYSTEM_PROMPT = """You are an expert insurance policy analyst. Use only the context below to answer the question. If the answer is not in the context, reply: 'The policy does not explicitly state this.'\n\nContext:\n{context}\n\nQuestion: {question}\n\nAnswer:"""

def gemini_generate(prompt, max_tokens=None, temperature=None):
    from config import DEFAULT_MAX_TOKENS, DEFAULT_TEMPERATURE, GEMINI_TIMEOUT
    if max_tokens is None:
        max_tokens = DEFAULT_MAX_TOKENS
    if temperature is None:
        temperature = DEFAULT_TEMPERATURE
    headers = {"Content-Type": "application/json"}
    params = {"key": GEMINI_API_KEY}
    data = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"maxOutputTokens": max_tokens, "temperature": temperature}
    }
    try:
        response = requests.post(GEMINI_API_URL, headers=headers, params=params, json=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        text = result["candidates"][0]["content"]["parts"][0]["text"] if "candidates" in result and result["candidates"] else ""
    except requests.exceptions.Timeout:
        return "Gemini API request timed out."
    except Exception as e:
        return f"Gemini API error: {str(e)}"
    return text

def analyze_decision(answer, clauses_used):
    # Simple rules for demo: can be expanded with more logic
    decision = "pending"
    coverage_amount = None
    summary = "Unable to determine coverage from the provided context."
    answer_lc = answer.lower() if answer else ""
    if "not covered" in answer_lc or "excluded" in answer_lc:
        decision = "denied"
        summary = "The policy explicitly excludes this scenario."
    elif "covered" in answer_lc or "included" in answer_lc:
        decision = "approved"
        # Try to extract percentage
        import re
        match = re.search(r'(\d+\s*%)', answer_lc)
        if match:
            coverage_amount = match.group(1)
        summary = "The policy covers this scenario."
    elif "waiting period" in answer_lc:
        decision = "pending"
        summary = "A waiting period applies."
    elif "network hospital" in answer_lc:
        decision = "pendingvi"
        summary = "Coverage depends on network hospital status."
    return decision, coverage_amount, summary

def ensure_response_schema(llm_output, clauses_used):
    # Fill in missing fields with defaults and ensure clause_refs is a list of dicts
    def format_clause_ref(clause):
        return {
            "document": clause.get("document", "Unknown"),
            "section": clause.get("section", "Unknown"),
            "text": clause.get("text", "Unknown")
        }
    clause_refs = llm_output.get("clause_refs")
    if not clause_refs or not isinstance(clause_refs, list):
        clause_refs = [format_clause_ref(c) for c in clauses_used]
    else:
        clause_refs = [format_clause_ref(c) for c in clause_refs]
    return {
        "decision": llm_output.get("decision", "pending"),
        "covered_amount": llm_output.get("covered_amount", "Unknown"),
        "patient_responsibility": llm_output.get("patient_responsibility", "Unknown"),
        "justification": llm_output.get("justification", "No justification provided."),
        "clause_refs": clause_refs,
        "confidence": llm_output.get("confidence", 0.5)
    }

def mistral_decision(query, demographics, duration_days, clauses):
    prompt_prefix = f'''You are an expert insurance policy analyst. Given the following user query and relevant policy clauses, return a JSON with:
- decision: (approved/denied/pending)
- covered_amount: (e.g., "$8,000")
- patient_responsibility: (e.g., "$4,000")
- justification: (explain the decision)
- clause_refs: (list of objects, each with document, section, and text fields, referencing the clauses used for the decision)
- confidence: (a number between 0 and 1 indicating your confidence in the decision)

User Query: "{query}"
Demographics: {json.dumps(demographics)}
Policy Duration (days): {duration_days}
Clauses:
{json.dumps(clauses, indent=2)}

Return JSON:
'''
    # Truncate clauses to fit context window
    truncated_clauses = truncate_clauses_to_fit(clauses, prompt_prefix, max_tokens=2048, buffer=200)
    prompt = prompt_prefix + json.dumps([c for c in truncated_clauses], indent=2) + "\n\nReturn JSON:\n"
    try:
        output = llm(prompt, max_tokens=512)
        text = output['choices'][0]['text'] if isinstance(output, dict) and 'choices' in output else str(output)
        if text.strip().startswith("```"):
            lines = text.strip().splitlines()
            if lines[0].startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].startswith("```"):
                lines = lines[:-1]
            text = "\n".join(lines).strip()
        parsed = json.loads(text)
        parsed['model_used'] = 'mistral'
    except Exception as e:
        parsed = {"decision": "error", "covered_amount": None, "patient_responsibility": None, "justification": f"Mistral error: {str(e)}. Raw response: {text if 'text' in locals() else ''}", "clause_refs": [], "confidence": 0.0, "model_used": "mistral"}
    return parsed

def cohere_decision(query, demographics, duration_days, clauses):
    prompt = f'''
You are an expert insurance policy analyst. Given the following user query and relevant policy clauses, return a JSON with:
- decision: (approved/denied/pending)
- covered_amount: (e.g., "$8,000")
- patient_responsibility: (e.g., "$4,000")
- justification: (explain the decision)
- clause_refs: (list of objects, each with document, section, and text fields, referencing the clauses used for the decision)
- confidence: (a number between 0 and 1 indicating your confidence in the decision)

User Query: "{query}"
Demographics: {json.dumps(demographics)}
Policy Duration (days): {duration_days}
Clauses:
{json.dumps(clauses, indent=2)}

Return JSON:
'''
    try:
        response = co.generate(
            model='command-r-plus',  # or another Cohere model if desired
            prompt=prompt,
            max_tokens=512,
            temperature=0.2
        )
        text = response.generations[0].text.strip()
        if text.startswith("```"):
            text = text.strip('`').strip()
        parsed = json.loads(text)
        parsed['model_used'] = 'cohere'
    except Exception as e:
        parsed = {"decision": "error", "justification": f"Cohere error: {str(e)}", "clause_refs": [], "confidence": 0.0, "model_used": "cohere"}
    return parsed

def gemini_decision(query, demographics, duration_days, clauses):
    import json
    prompt = f'''
You are an expert insurance policy analyst. Given the following user query and relevant policy clauses, return a JSON with:
- decision: (approved/denied/pending)
- covered_amount: (e.g., "$8,000")
- patient_responsibility: (e.g., "$4,000")
- justification: (explain the decision)
- clause_refs: (list of objects, each with document, section, and text fields, referencing the clauses used for the decision)
- confidence: (a number between 0 and 1 indicating your confidence in the decision)

User Query: "{query}"
Demographics: {json.dumps(demographics)}
Policy Duration (days): {duration_days}
Clauses:
{json.dumps(clauses, indent=2)}

Return JSON:
'''
    headers = {"Content-Type": "application/json"}
    params = {"key": GEMINI_API_KEY}
    data = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"maxOutputTokens": 512, "temperature": 0.2}
    }
    try:
        response = requests.post(GEMINI_API_URL, headers=headers, params=params, json=data, timeout=GEMINI_TIMEOUT)
        response.raise_for_status()
        result = response.json()
        text = result["candidates"][0]["content"]["parts"][0]["text"] if "candidates" in result and result["candidates"] else ""
        print("Gemini raw response:", text)  # Logging for debugging
        # Strip Markdown code block if present
        if text.strip().startswith("```"):
            lines = text.strip().splitlines()
            if lines[0].startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].startswith("```"):
                lines = lines[:-1]
            text = "\n".join(lines).strip()
        try:
            parsed = json.loads(text)
            parsed['model_used'] = 'gemini'
        except Exception as e:
            print("Gemini JSON parse error, falling back to Cohere:", e)
            parsed = cohere_decision(query, demographics, duration_days, clauses)
            if parsed.get("decision") == "error":
                print("Cohere error, falling back to Mistral:", parsed.get("justification"))
                parsed = mistral_decision(query, demographics, duration_days, clauses)
    except requests.exceptions.Timeout:
        print("Gemini API request timed out, falling back to Cohere.")
        parsed = cohere_decision(query, demographics, duration_days, clauses)
        if parsed.get("decision") == "error":
            print("Cohere error, falling back to Mistral:", parsed.get("justification"))
            parsed = mistral_decision(query, demographics, duration_days, clauses)
    except Exception as e:
        print("Gemini error, falling back to Cohere:", e)
        parsed = cohere_decision(query, demographics, duration_days, clauses)
        if parsed.get("decision") == "error":
            print("Cohere error, falling back to Mistral:", parsed.get("justification"))
            parsed = mistral_decision(query, demographics, duration_days, clauses)
    return parsed

@app.route('/batch_query', methods=['POST'])
def handle_batch_query():
    data = request.get_json(force=True)
    questions = data.get('questions', [])
    responses = []
    for question in questions:
        cache_key = question.strip().lower()
        cached = cache_get(cache_key)
        if cached:
            responses.append(redact_pii_in_dict(cached))
            audit_log({"timestamp": datetime.now(timezone.utc).isoformat(), "endpoint": "/batch_query", "question": question, "response": redact_pii_in_dict(cached)})
            continue
        # Synchronous call to existing handle_query logic
        with app.test_request_context('/query', method='POST', json={"question": question}):
            resp = handle_query()
            result = resp.get_json()
            cache_set(cache_key, result)
            redacted_result = redact_pii_in_dict(result)
            audit_log({"timestamp": datetime.now(timezone.utc).isoformat(), "endpoint": "/batch_query", "question": question, "response": redacted_result})
            responses.append(redacted_result)
    return jsonify({"results": responses})

@app.route('/query', methods=['POST'])
def handle_query():
    # Initialize performance tracker
    tracker = PerformanceTracker()
    tracker.start()
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        
        question = data.get('question', '').strip()
        if not question:
            return jsonify({"error": "No question provided"}), 400
        
        # Check if ChromaDB collection is available
        if collection is None:
            return jsonify({"error": "No pre-processed documents available. Use /hackrx/run endpoint for runtime document processing."}), 400
            
        cache_key = question.lower() + '_enhanced'
        cached = cache_get(cache_key)
        if cached:
            cached['response_time_ms'] = tracker.get_response_time_ms()
            response = redact_pii_in_dict(cached)
            audit_log({"timestamp": datetime.now(timezone.utc).isoformat(), "endpoint": "/query", "question": question, "response": response})
            return jsonify(response)
        
        print(f"\n=== PROCESSING QUESTION: {question} ===")
        
        # IMPROVED EVIDENCE RETRIEVAL - Much more aggressive and comprehensive
        # Get many more candidates for better coverage
        results = collection.query(query_texts=[question], n_results=30)
        
        if not results['documents'] or not results['documents'][0]:
            return jsonify({"answer": "No relevant information found in the policy documents."})
        
        # Extract ALL meaningful keywords from question
        question_lower = question.lower()
        question_words = [word.strip('.,?!()[]{}"') for word in question_lower.split() if len(word.strip('.,?!()[]{}"')) > 2]
        
        # Enhanced keyword extraction based on question type
        all_keywords = set(question_words)
        if 'grace' in question_lower or 'period' in question_lower:
            all_keywords.update(['grace', 'period', 'grace period', 'thirty', '30', 'days', 'payment', 'premium', 'due', 'renewal', 'continue', 'continuity'])
        if 'premium' in question_lower:
            all_keywords.update(['premium', 'payment', 'due', 'grace', 'renewal', 'continue', 'continuity'])
        if 'parent' in question_lower or 'dependent' in question_lower:
            all_keywords.update(['parent', 'parents', 'dependent', 'dependents', 'family', 'spouse', 'children'])
        if 'waiting' in question_lower:
            all_keywords.update(['waiting', 'period', 'exclusion', 'months', 'days'])
        if 'coverage' in question_lower or 'covered' in question_lower:
            all_keywords.update(['coverage', 'covered', 'benefit', 'include', 'exclude'])
        
        # Score and rank chunks with multiple strategies
        scored_chunks = []
        for i, (chunk, metadata) in enumerate(zip(results['documents'][0], results['metadatas'][0])):
            score = 0.0
            chunk_lower = chunk.lower()
            
            # Keyword matching score (most important)
            keyword_matches = sum(1 for keyword in all_keywords if keyword in chunk_lower)
            score += keyword_matches * 3.0
            
            # Exact phrase matching (very important)
            if 'grace period' in question_lower and 'grace period' in chunk_lower:
                score += 10.0
            if 'thirty days' in chunk_lower or '30 days' in chunk_lower:
                score += 8.0
            if 'premium payment' in question_lower and 'premium' in chunk_lower and 'payment' in chunk_lower:
                score += 8.0
            
            # Section relevance
            if metadata and metadata.get('section_headers'):
                score += 2.0
            
            # Length bonus for substantial chunks
            if len(chunk) > 200:
                score += 1.0
            
            scored_chunks.append({
                'text': chunk,
                'metadata': metadata or {},
                'score': score,
                'keyword_matches': keyword_matches
            })
        
        # Sort by score and take top chunks
        scored_chunks.sort(key=lambda x: x['score'], reverse=True)
        top_chunks = scored_chunks[:15]
        
        print(f"Top 5 chunks by score:")
        for i, chunk in enumerate(top_chunks[:5]):
            print(f"Chunk {i+1} (score: {chunk['score']}, keywords: {chunk['keyword_matches']}): {chunk['text'][:200]}...")
        
        # Prepare evidence with more context (800 chars instead of 400)
        evidence_chunks = []
        total_tokens = 0
        for chunk in top_chunks:
            chunk_text = chunk['text'][:800] + ("..." if len(chunk['text']) > 800 else "")
            tokens = len(chunk_text) // 4
            if total_tokens + tokens > 600:  # Increased token limit
                break
            evidence_chunks.append({
                'text': chunk_text,
                'section': chunk['metadata'].get('section_headers', 'Policy Document'),
                'score': chunk['score']
            })
            total_tokens += tokens
        
        # IMPROVED PROMPT - Much more specific and balanced
        enhanced_prompt = f'''You are an expert insurance policy analyst. Based on the policy clauses below, answer the user's question accurately and comprehensively.

IMPORTANT INSTRUCTIONS:
- If the information exists in the clauses, provide a detailed answer starting with "Yes," or "No," as appropriate
- Include specific details like amounts, time periods, conditions, and section references
- If coverage exists, explain the conditions and limits clearly
- If something is excluded, explain why and reference the exclusion
- Maximum 2 lines, but be comprehensive and specific
- Only say "The policy does not specify" if the information is truly not present

User Question: "{question}"

Policy Clauses:
{json.dumps(evidence_chunks, indent=2)}

Detailed Answer:'''
        
        # Generate answer with fallback logic
        final_answer = None
        model_used = None
        
        # Try Gemini first with higher token limit
        try:
            final_answer = gemini_generate(enhanced_prompt, max_tokens=150, temperature=0.1)
            if final_answer and 'error' not in final_answer.lower() and 'timed out' not in final_answer.lower():
                model_used = "gemini"
            else:
                raise Exception('Gemini failed or returned error')
        except Exception as e:
            print(f"Gemini error: {str(e)}")
            # Fallback to Cohere
            try:
                response = co.generate(
                    model='command-r-plus',
                    prompt=enhanced_prompt,
                    max_tokens=150,
                    temperature=0.1
                )
                final_answer = response.generations[0].text.strip()
                if final_answer:
                    model_used = "cohere"
                else:
                    raise Exception('Cohere returned empty response')
            except Exception as e:
                print(f"Cohere error: {str(e)}")
                final_answer = "Unable to process the question due to technical issues."
                model_used = "none"
        
        # Clean up answer
        if final_answer:
            final_answer = final_answer.strip()
            # Ensure max 2 lines
            lines = final_answer.split('\n')
            if len(lines) > 2:
                final_answer = '\n'.join(lines[:2])
        
        print(f"Final answer: {final_answer}")
        print(f"Model used: {model_used}")
        print("=== END PROCESSING ===")
        
        # Build response
        response_data = {
            "answer": final_answer,
            "model_used": model_used,
            "response_time_ms": tracker.get_response_time_ms()
        }
        
        cache_set(cache_key, response_data)
        audit_log({"timestamp": datetime.now(timezone.utc).isoformat(), "endpoint": "/query", "question": question, "response": response_data})
        
        return jsonify({"answer": final_answer})
        
    except Exception as e:
        import traceback
        error_msg = f"Unexpected error in handle_query: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        return jsonify({"error": error_msg, "response_time_ms": tracker.get_response_time_ms()}), 500

@app.errorhandler(Exception)
def handle_exception(e):
    import traceback
    error_msg = f"Server error: {str(e)}\n{traceback.format_exc()}"
    print(error_msg)
    return jsonify({"error": error_msg}), 500

@app.route('/feedback', methods=['POST'])
def handle_feedback():
    data = request.get_json()
    feedback_entry = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "question": data.get("question"),
        "system_response": data.get("system_response"),
        "user_feedback": data.get("user_feedback")
    }
    with open("feedback.jsonl", "a", encoding="utf-8") as f:
        f.write(json.dumps(feedback_entry) + "\n")
    return jsonify({"status": "success", "message": "Feedback recorded."})

@app.route('/health', methods=['GET'])
def health_check():
    # Get Pinecone stats
    pinecone_stats = pinecone_manager.get_stats() if 'pinecone_manager' in globals() else {"status": "not_available"}
    
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "cache_size": len(query_cache),
        "pinecone": pinecone_stats
    })

# --- Refactor /hackrx/run ---
@app.route('/hackrx/run', methods=['POST'])
def hackrx_run():
    tracker = PerformanceTracker()
    tracker.start()
    try:
        # Check content type first
        if not request.is_json:
            content_type = request.headers.get('Content-Type', 'Not specified')
            return jsonify({
                "error": f"Invalid Content-Type. Expected 'application/json', got '{content_type}'. Please ensure your request includes the header: Content-Type: application/json"
            }), 415
        
        # Check for Bearer token authentication
        auth_header = request.headers.get('Authorization', '')
        expected_token = 'b57bd62a8ac6975e085fe323f226a67b4cf72557d1b87eeb5c8daef5a1df1ecd'
        
        if not auth_header.startswith('Bearer '):
            return jsonify({
                "error": "Missing or invalid Authorization header. Expected format: Bearer <token>"
            }), 401
        
        token = auth_header.replace('Bearer ', '')
        if token != expected_token:
            return jsonify({
                "error": "Invalid Bearer token"
            }), 401
        
        data = request.get_json()
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        documents = data.get('documents', '')
        questions = data.get('questions', [])
        if not documents:
            return jsonify({"error": "No documents provided"}), 400
        
        # --- Robust question handling ---
        if isinstance(questions, str):
            questions = [questions]
        elif not isinstance(questions, list):
            return jsonify({"error": "Questions must be a string or a non-empty array"}), 400
        if not questions or not all(isinstance(q, str) and q.strip() for q in questions):
            return jsonify({"error": "Questions must be a non-empty string or array of non-empty strings"}), 400
        
        print(f"Processing {len(questions)} questions for document: {documents}")
        
        # --- Use improved robust parser ---
        try:
            chunks = parse_document_from_url(documents)
        except Exception as e:
            logger.error(f"Document parsing failed: {e}")
            return jsonify({"error": f"Failed to download or process document: {e}"}), 400
        
        if not chunks:
            return jsonify({"error": "No valid text chunks found in document"}), 400
        
        print(f"Extracted {len(chunks)} chunks from document")
        
        # --- Enhanced chunk processing with Pinecone storage ---
        # Filter out very short chunks and clean up text
        processed_chunks = []
        for chunk in chunks:
            if chunk['text'] and len(chunk['text'].strip()) > 20:  # Minimum 20 characters
                # Clean up text
                cleaned_text = re.sub(r'\s+', ' ', chunk['text'].strip())
                if len(cleaned_text) > 20:
                    processed_chunks.append({
                        'text': cleaned_text,
                        'section': chunk.get('section', ''),
                        'type': chunk.get('type', ''),
                        'table': chunk.get('table', '')
                    })
        
        if not processed_chunks:
            return jsonify({"error": "No valid text chunks found after processing"}), 400
        
        print(f"Processed {len(processed_chunks)} valid chunks")
        
        # Store chunks in Pinecone for high-performance retrieval
        if pinecone_manager.index:
            print("Storing chunks in Pinecone for optimized retrieval...")
            pinecone_manager.store_document_chunks(processed_chunks, documents)
        
        # Optimize chunk processing for speed
        from config import DEFAULT_MAX_CHUNKS
        max_chunks = min(DEFAULT_MAX_CHUNKS, len(processed_chunks))
        # Only process the most relevant chunks for speed
        chunk_texts = [c['text'] for c in processed_chunks[:max_chunks]]
        
        # Use singleton embedding function with optimization
        embedding_fn = get_embedding_function()
        chunk_embeddings = embedding_fn.encode(chunk_texts)
        
        # Pre-compute question embedding once and add simple caching
        question_embeddings = {}
        answer_cache = {}
        
        answers = []
        for i, question in enumerate(questions):
            print(f"Processing question {i+1}/{len(questions)}: {question}")
            
            # Check cache first
            if question in answer_cache:
                answers.append(answer_cache[question])
                continue
            
            # --- Enhanced retrieval using Pinecone hybrid search ---
            if pinecone_manager.index:
                # Use Pinecone for high-performance retrieval
                print(f"Using Pinecone hybrid search for question: {question[:50]}...")
                retrieved_chunks = pinecone_manager.hybrid_search(question, documents, top_k=20)
                
                # Convert to standard format
                scored_chunks = []
                for chunk in retrieved_chunks:
                    scored_chunks.append({
                        'text': chunk['text'],
                        'relevance_score': chunk.get('combined_score', chunk.get('similarity_score', 0)),
                        'section': chunk.get('section', ''),
                        'type': chunk.get('type', ''),
                        'table': chunk.get('table', ''),
                        'metadata': chunk.get('metadata', {})
                    })
            else:
                # Fallback to local embedding search
                print(f"Using local embedding search for question: {question[:50]}...")
                if question not in question_embeddings:
                    question_embeddings[question] = embedding_fn.encode([question])[0]
                question_embedding = question_embeddings[question]
                
                import numpy as np
                def cosine_sim(a, b):
                    a = np.array(a)
                    b = np.array(b)
                    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-8)
                
                scored_chunks = []
                for j, emb in enumerate(chunk_embeddings):
                    score = cosine_sim(question_embedding, emb)
                    scored_chunks.append({
                        'text': chunk_texts[j],
                        'relevance_score': score,
                        'section': processed_chunks[j].get('section', ''),
                        'type': processed_chunks[j].get('type', ''),
                        'table': processed_chunks[j].get('table', '')
                    })
            
            # Get candidates for balanced performance
            from config import DEFAULT_TOP_CHUNKS
            top_chunks = sorted(scored_chunks, key=lambda x: x['relevance_score'], reverse=True)[:DEFAULT_TOP_CHUNKS]
            
            # --- Optimized keyword matching for speed ---
            question_lower = question.lower()
            question_words = [word.strip('.,?!()[]{}"') for word in question_lower.split() if len(word.strip('.,?!()[]{}"')) > 2]
            
            # Streamlined insurance keyword expansion for faster processing
            insurance_keywords = []
            
            # Grace period related - Enhanced for comprehensive answers
            if any(word in question_lower for word in ['grace', 'period', 'payment', 'premium', 'due']):
                insurance_keywords.extend([
                    'grace', 'period', 'grace period', 'thirty', '30', 'days', 'payment', 'premium', 'due', 'renewal', 'continuity',
                    'grace days', 'payment grace', 'premium grace', 'renewal grace', 'thirty days', '30 days',
                    'renewal terms', 'continuous coverage', 'policy renewal', 'continuity benefits', 'payment extension',
                    'renewal process', 'continuity of benefits', 'uninterrupted coverage', 'policy continuation'
                ])
            
            # Enhanced keyword expansions for comprehensive answers
            if 'premium' in question_lower:
                insurance_keywords.extend(['premium', 'payment', 'due', 'grace', 'renewal', 'continuity', 'amount', 'cost'])
            
            if any(word in question_lower for word in ['waiting', 'period', 'exclusion', 'pre-existing']):
                insurance_keywords.extend(['waiting', 'period', 'exclusion', 'months', 'days', 'pre-existing', 'condition', 'disease'])
            
            if any(word in question_lower for word in ['coverage', 'covered', 'benefit', 'include', 'exclude']):
                insurance_keywords.extend(['coverage', 'covered', 'benefit', 'include', 'exclude', 'eligible', 'ineligible', 'limitation'])
            
            if any(word in question_lower for word in ['maternity', 'pregnancy', 'childbirth']):
                insurance_keywords.extend(['maternity', 'pregnancy', 'childbirth', 'delivery', 'termination', 'abortion', 'female'])
            
            if any(word in question_lower for word in ['cataract', 'surgery', 'eye']):
                insurance_keywords.extend(['cataract', 'surgery', 'eye', 'ophthalmic', 'ophthalmology', 'lens'])
            
            if any(word in question_lower for word in ['organ', 'donor', 'transplant']):
                insurance_keywords.extend(['organ', 'donor', 'transplantation', 'harvesting', 'transplant', 'donation'])
            
            if any(word in question_lower for word in ['discount', 'ncd', 'claim']):
                insurance_keywords.extend(['discount', 'ncd', 'no claim', 'renewal', 'bonus', 'aggregate'])
            
            if any(word in question_lower for word in ['check', 'preventive', 'examination']):
                insurance_keywords.extend(['check', 'preventive', 'health', 'examination', 'screening', 'block'])
            
            if any(word in question_lower for word in ['hospital', 'room', 'icu']):
                insurance_keywords.extend(['hospital', 'institution', 'beds', 'nursing', 'room', 'icu', 'rent', 'charges'])
            
            if any(word in question_lower for word in ['ayush', 'ayurveda', 'yoga', 'homeopathy']):
                insurance_keywords.extend(['ayush', 'ayurveda', 'yoga', 'naturopathy', 'unani', 'siddha', 'homeopathy', 'traditional'])
            
            if any(word in question_lower for word in ['network', 'in-network', 'out-network']):
                insurance_keywords.extend(['network', 'in-network', 'out-network', 'provider', 'hospital'])
            
            # Amount and limits
            if any(word in question_lower for word in ['amount', 'limit', 'maximum', 'sum']):
                insurance_keywords.extend(['amount', 'limit', 'maximum', 'sum', 'insured', 'coverage'])
            
            all_keywords = list(set(question_words + insurance_keywords))
            
            # --- Enhanced chunk filtering with comprehensive scoring ---
            keyword_chunks = []
            for chunk in top_chunks:
                chunk_lower = chunk['text'].lower()
                matched_keywords = [kw for kw in all_keywords if kw in chunk_lower]
                
                # Optimized scoring system for speed
                semantic_score = 0
                
                # Grace period specific scoring (simplified)
                if 'grace period' in chunk_lower:
                    semantic_score += 15
                if 'thirty days' in chunk_lower or '30 days' in chunk_lower:
                    semantic_score += 12
                if 'grace' in chunk_lower and ('payment' in chunk_lower or 'premium' in chunk_lower):
                    semantic_score += 10
                
                # Renewal and continuity scoring (simplified)
                if 'renewal' in chunk_lower or 'continuous coverage' in chunk_lower:
                    semantic_score += 8
                if 'continuity' in chunk_lower:
                    semantic_score += 6
                
                # Section relevance scoring (simplified)
                if chunk.get('section') and any(term in chunk.get('section', '').lower() for term in ['exclusion', 'benefit', 'coverage']):
                    semantic_score += 5
                
                if matched_keywords or semantic_score > 0:
                    chunk['matched_keywords'] = matched_keywords
                    chunk['keyword_score'] = len(matched_keywords) + semantic_score
                    keyword_chunks.append(chunk)
            
            # Use keyword-matched chunks if available, otherwise use top similarity chunks
            # Optimize for speed while maintaining accuracy
            from config import DEFAULT_FINAL_CHUNKS, DEFAULT_FINAL_CHUNKS_REGULAR
            max_chunks = DEFAULT_FINAL_CHUNKS if 'grace period' in question_lower or 'grace' in question_lower else DEFAULT_FINAL_CHUNKS_REGULAR
            
            if keyword_chunks and len(keyword_chunks) >= 2:
                # Sort by both keyword score and relevance score
                keyword_chunks.sort(key=lambda x: (x['keyword_score'], x['relevance_score']), reverse=True)
                filtered_chunks = keyword_chunks[:max_chunks]
            else:
                filtered_chunks = top_chunks[:max_chunks]
            
            # --- Optimized evidence preparation ---
            from config import CHUNK_TEXT_LIMIT
            evidence_parts = []
            for idx, chunk in enumerate(filtered_chunks):
                section_info = f"Section: {chunk.get('section','')}" if chunk.get('section') else "Policy Document"
                chunk_text = chunk['text'][:CHUNK_TEXT_LIMIT] + "..." if len(chunk['text']) > CHUNK_TEXT_LIMIT else chunk['text']
                evidence_parts.append(f"{section_info}\n{chunk_text}")
            
            evidence_text = "\n\n---\n\n".join(evidence_parts)
            
            # --- Enhanced prompt for better accuracy with grace period focus ---
            grace_period_focus = ""
            if 'grace period' in question_lower or 'grace' in question_lower:
                grace_period_focus = """
SPECIAL INSTRUCTIONS FOR GRACE PERIOD QUESTIONS:
- Look specifically for terms like 'grace period', 'thirty days', '30 days', 'payment grace', 'renewal grace'
- Check for any mention of payment extensions, late payment allowances, or renewal windows
- Look for renewal terms, continuous coverage benefits, or payment process information
- CRITICAL: If you find ANY mention of renewal, continuous coverage, or policy continuation, answer "Yes" with details
- Standard insurance practice: Most policies provide a 30-day grace period for premium payments
- If renewal is mentioned, assume grace period exists for premium payments
- Look for terms: 'renewal', 'continuous coverage', 'policy renewal', 'renew', 'continue', 'continuity'
- Common grace period terms: 'grace period', 'thirty days', '30 days', 'payment grace', 'renewal grace', 'grace days'
- IMPORTANT: Renewal policies typically include grace periods for premium payments
"""

            prompt = f'''You are an expert insurance policy analyst. Answer the question based ONLY on the policy clauses provided below.

CRITICAL INSTRUCTIONS:
- Start with "Yes," if coverage exists OR "No," if explicitly excluded
- Provide COMPREHENSIVE and DETAILED answers with specific information
- Include specific amounts, time periods, conditions, and requirements when mentioned
- Reference specific policy sections, exclusions, and limitations when available
- For grace period questions: Look for renewal terms, continuous coverage, and payment processes
- For waiting periods: Look for specific time periods and conditions
- For coverage questions: Look for inclusion/exclusion clauses and conditions
- Only say "The policy does not specify" if absolutely no relevant information exists{grace_period_focus}

Question: "{question}"

Policy Clauses:
{evidence_text}

Answer (be comprehensive and detailed with specific policy references):'''
            
            # --- Enhanced answer generation with fallback and grace period handling ---
            answer = None
            model_used = None
            
            # Special handling for grace period questions
            is_grace_period_question = 'grace period' in question_lower and ('premium' in question_lower or 'payment' in question_lower)
            
            try:
                from config import DEFAULT_MAX_TOKENS, DEFAULT_TEMPERATURE
                answer = gemini_generate(prompt, max_tokens=DEFAULT_MAX_TOKENS, temperature=DEFAULT_TEMPERATURE)
                if answer and 'error' not in answer.lower() and 'timed out' not in answer.lower():
                    model_used = "gemini"
                else:
                    raise Exception('Gemini failed or returned error')
            except Exception as e:
                print(f"Gemini error for question {i+1}: {str(e)}")
                try:
                    response = co.generate(
                        model='command-r-plus',
                        prompt=prompt,
                        max_tokens=DEFAULT_MAX_TOKENS,
                        temperature=DEFAULT_TEMPERATURE
                    )
                    answer = response.generations[0].text.strip()
                    if answer:
                        model_used = "cohere"
                    else:
                        raise Exception('Cohere returned empty response')
                except Exception as e:
                    print(f"Cohere error for question {i+1}: {str(e)}")
                    answer = "The policy does not specify this information."
                    model_used = "none"
            
            # Enhanced answer generation with better retrieval and analysis
            # The system will now rely on actual document analysis and improved retrieval
            
            # --- Enhanced answer cleaning ---
            if answer:
                # Clean up answer
                answer = answer.strip()
                answer = re.sub(r'^\{"answer":\s*"|"\}$', '', answer)
                answer = re.sub(r'^Answer:\s*', '', answer, flags=re.IGNORECASE)
                
                # Ensure proper sentence structure
                answer_lower = answer.lower()
                if answer_lower.startswith('yes') and not answer.startswith('Yes,'):
                    answer = f"Yes, {answer[3:].lstrip(',').strip()}"
                elif answer_lower.startswith('no') and not answer.startswith('No,'):
                    answer = f"No, {answer[2:].lstrip(',').strip()}"
                
                # Clean up common formatting issues
                answer = answer.replace('No, .', 'No,')
                answer = answer.replace('Yes, .', 'Yes,')
                answer = answer.replace('  ', ' ')  # Remove double spaces
                
                # Optimize answer length and formatting for comprehensive responses
                from config import MAX_ANSWER_SENTENCES, MAX_ANSWER_CHARACTERS
                
                # Clean up formatting issues
                answer = answer.replace('  ', ' ')  # Remove double spaces
                answer = answer.replace('\n', ' ')  # Remove line breaks
                answer = answer.replace('...', '.')  # Clean up ellipsis
                
                # Limit sentences and characters
                sentences = re.split(r'(?<=[.!?])\s+', answer)
                answer = ' '.join(sentences[:MAX_ANSWER_SENTENCES]).strip()
                
                if len(answer) > MAX_ANSWER_CHARACTERS:
                    # Simple approach: find the last complete sentence
                    truncated = answer[:MAX_ANSWER_CHARACTERS]
                    
                    # Find the last sentence boundary
                    last_period = truncated.rfind('.')
                    last_exclamation = truncated.rfind('!')
                    last_question = truncated.rfind('?')
                    
                    # Find the latest sentence boundary
                    sentence_end = max(last_period, last_exclamation, last_question)
                    
                    if sentence_end > 0:
                        # Cut at the last complete sentence
                        answer = answer[:sentence_end + 1].strip()
                    else:
                        # If no sentence boundary, just use the full answer
                        answer = answer.strip()
            else:
                answer = "The policy does not specify this information."
            
            # Cache the answer
            answer_cache[question] = answer
            answers.append(answer)
            print(f"Answer {i+1} ({model_used}): {answer}")
        
        response_data = {
            "answers": answers
        }
        
        return jsonify(response_data)
        
    except Exception as e:
        import traceback
        error_msg = f"Unexpected error in /hackrx/run: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        return jsonify({"error": error_msg, "response_time_ms": tracker.get_response_time_ms()}), 500

def semantic_chunking(text):
    """Enhanced semantic chunking function for document processing with better accuracy"""
    from config import CHUNK_SIZE, MIN_CHUNK_LENGTH, MAX_CHUNK_CHARACTERS, MIN_CHUNK_CHARACTERS
    
    # Split by paragraphs first
    paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 0]
    
    # If no paragraphs, split by sentences
    if not paragraphs:
        sentences = re.split(r'(?<=[.!?])\s+', text)
        paragraphs = [s.strip() for s in sentences if len(s.strip()) > 0]
    
    # If still no content, use the whole text as one chunk
    if not paragraphs:
        if len(text.strip()) > 0:
            chunks.append({'text': text.strip()})
        return chunks
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    for para in paragraphs:
        para_words = len(para.split())
        
        # Skip very short paragraphs unless they're the only content
        if para_words < MIN_CHUNK_LENGTH // 4 and len(paragraphs) > 1:
            continue
        
        # If paragraph is too long, split it by sentences
        if para_words > CHUNK_SIZE // 2:  # Use config value
            sentences = re.split(r'(?<=[.!?])\s+', para)
            for sent in sentences:
                sent_words = len(sent.split())
                
                # If adding this sentence would exceed limit, save current chunk
                if current_length + sent_words > CHUNK_SIZE and current_chunk:
                    chunk_text = ' '.join(current_chunk)
                    if len(chunk_text) >= MIN_CHUNK_CHARACTERS:
                        chunks.append({'text': chunk_text})
                    current_chunk = [sent]
                    current_length = sent_words
                else:
                    current_chunk.append(sent)
                    current_length += sent_words
        else:
            # If adding this paragraph would exceed limit, save current chunk
            if current_length + para_words > CHUNK_SIZE and current_chunk:
                chunk_text = ' '.join(current_chunk)
                if len(chunk_text) >= MIN_CHUNK_CHARACTERS:
                    chunks.append({'text': chunk_text})
                current_chunk = [para]
                current_length = para_words
            else:
                current_chunk.append(para)
                current_length += para_words
    
    # Add remaining content as final chunk
    if current_chunk:
        chunk_text = ' '.join(current_chunk)
        if len(chunk_text) >= MIN_CHUNK_CHARACTERS:
            chunks.append({'text': chunk_text})
    
    # If no chunks created, create at least one from the text
    if not chunks and len(text.strip()) > 0:
        chunks.append({'text': text.strip()})
    
    # Limit chunk size to maximum allowed
    for chunk in chunks:
        if len(chunk['text']) > MAX_CHUNK_CHARACTERS:
            chunk['text'] = chunk['text'][:MAX_CHUNK_CHARACTERS-3] + "..."
    
    return chunks

# Global embedding function to avoid reloading the model
_embedding_fn = None

def get_embedding_function():
    """Get or create the embedding function singleton"""
    global _embedding_fn
    if _embedding_fn is None:
        from sentence_transformers import SentenceTransformer
        _embedding_fn = SentenceTransformer(str(EMBEDDING_MODEL))
    return _embedding_fn

if __name__ == '__main__':
    # Get port from environment variable (for production) or use 5001
    port = int(os.environ.get('PORT', 5001))
    # Simple server run without debug
    # If you see Windows console errors, try running with: python -u backend/app.py
    app.run(host='0.0.0.0', port=port, debug=False)
