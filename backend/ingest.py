import pdfplumber


from sentence_transformers import SentenceTransformer
from config import *
from pathlib import Path
import re
from typing import List, Dict
import uuid
import pytesseract
from pdf2image import convert_from_path
import camelot
from datetime import datetime
import docx
from docx import Document
import extract_msg
import email
from email import policy
from email.parser import BytesParser
import json
import requests
from config import GEMINI_API_KEY, GEMINI_API_URL

class EnhancedDocumentProcessor:
    def __init__(self):
        self.embedding_model = SentenceTransformer(str(EMBEDDING_MODEL))
        

    def clean_text(self, text: str) -> str:
        """Enhanced text cleaning"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        # Remove page numbers and headers
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        # Clean up special characters
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}]', ' ', text)
        return text

    def extract_insurance_section_headers(self, text: str) -> List[str]:
        """Enhanced section header extraction for insurance documents"""
        # Insurance-specific patterns
        patterns = [
            r'^(?:Section|Clause|Article)\s+\d+[\.\d]*\s*[:\-]?\s*[A-Z][A-Za-z\s]+$',
            r'^(?:Coverage|Benefits|Exclusions|Limitations|Terms|Conditions)\s*[:\-]?\s*[A-Z][A-Za-z\s]+$',
            r'^(?:Policy|Insurance|Premium|Claim|Deductible|Co-pay)\s+[A-Z][A-Za-z\s]+$',
            r'^\d+\.\s+[A-Z][A-Za-z\s]+$',
            r'^[A-Z][A-Z\s]{3,}$',  # All caps headers
        ]
        
        headers = []
        for pattern in patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            headers.extend(matches)
        
        return list(set(headers))  # Remove duplicates

    def extract_insurance_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract insurance-specific entities using LLM"""
        prompt = f"""Extract insurance-specific entities from this text. Return as JSON:
{{
  "policy_types": ["health", "life", "auto"],
  "coverage_areas": ["surgery", "medication", "hospitalization"],
  "exclusions": ["pre-existing", "cosmetic"],
  "limits": ["$5000", "80%"],
  "waiting_periods": ["30 days", "6 months"]
}}

Text: {text[:1000]}"""

        try:
            headers = {"Content-Type": "application/json"}
            params = {"key": GEMINI_API_KEY}
            data = {
                "contents": [{"parts": [{"text": prompt}]}],
                "generationConfig": {"maxOutputTokens": 256, "temperature": 0.1}
            }
            response = requests.post(GEMINI_API_URL, headers=headers, params=params, json=data, timeout=10)
            response.raise_for_status()
            result = response.json()
            text = result["candidates"][0]["content"]["parts"][0]["text"] if "candidates" in result and result["candidates"] else ""
            
            # Clean JSON response
            if text.strip().startswith("```"):
                text = text.strip('`').strip()
            
            return json.loads(text)
        except Exception as e:
            print(f"LLM entity extraction error: {str(e)}")
            return {"policy_types": [], "coverage_areas": [], "exclusions": [], "limits": [], "waiting_periods": []}

    def extract_effective_dates(self, text: str) -> List[str]:
        """Enhanced date extraction"""
        date_patterns = [
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            r'\b\d{4}-\d{2}-\d{2}\b',
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b',
            r'\b(?:effective|valid|from|until)\s+(?:date|period)?\s*[:\-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b'
        ]
        dates = []
        for pat in date_patterns:
            dates += re.findall(pat, text, re.IGNORECASE)
        return list(set(dates))  # Remove duplicates

    def extract_footnotes(self, text: str) -> List[str]:
        """Enhanced footnote extraction"""
        footnote_patterns = [
            r'^(\*|\d+)\s+.+$',
            r'^Note\s*\d*[:\-]?\s*.+$',
            r'^Footnote\s*\d*[:\-]?\s*.+$'
        ]
        footnotes = []
        for pattern in footnote_patterns:
            footnotes.extend(re.findall(pattern, text, re.MULTILINE | re.IGNORECASE))
        return footnotes

    def extract_tables(self, pdf_path: Path) -> List[str]:
        """Enhanced table extraction"""
        tables = []
        try:
            c_tables = camelot.read_pdf(str(pdf_path), pages='all', flavor='stream')  # type: ignore
            for table in c_tables:
                # Clean table data
                table_text = table.df.to_string()
                # Remove empty rows and clean up
                lines = [line.strip() for line in table_text.split('\n') if line.strip()]
                if lines:
                    tables.append('\n'.join(lines))
        except Exception as e:
            print(f"Table extraction error: {str(e)}")
        return tables

    def extract_tables_docx(self, doc: Document) -> List[str]:
        """Enhanced DOCX table extraction"""
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                # Clean cell text
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:  # Only add non-empty rows
                    rows.append(' | '.join(cells))
            if rows:  # Only add non-empty tables
                tables.append('\n'.join(rows))
        return tables

    def ocr_pdf(self, pdf_path: Path) -> str:
        """Enhanced OCR with better error handling"""
        try:
            images = convert_from_path(str(pdf_path))
            ocr_text = []
            for img in images:
                text = pytesseract.image_to_string(img)
                if text.strip():  # Only add non-empty text
                    ocr_text.append(text)
            return '\n'.join(ocr_text)
        except Exception as e:
            print(f"OCR error: {str(e)}")
            return ''

    def semantic_chunking(self, text: str) -> List[Dict[str, any]]:
        """Section/paragraph-based chunking with quality scoring"""
        # Split by double newlines (paragraphs)
        paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 0]
        chunks = []
        for para in paragraphs:
            if len(para.split()) < MIN_CHUNK_LENGTH:
                continue
            # If paragraph is too long, split further by sentences
            if len(para.split()) > CHUNK_SIZE:
                sentences = re.split(r'(?<=[.!?])\s+', para)
                current_chunk = []
                current_len = 0
                for sent in sentences:
                    sent_len = len(sent.split())
                    if current_len + sent_len > CHUNK_SIZE and current_chunk:
                        chunk_text = ' '.join(current_chunk)
                        quality_score = self._calculate_chunk_quality(chunk_text)
                        chunks.append({
                            'text': chunk_text,
                            'quality_score': quality_score,
                            'word_count': len(chunk_text.split()),
                            'sentences': len(current_chunk)
                        })
                        current_chunk = []
                        current_len = 0
                    current_chunk.append(sent)
                    current_len += sent_len
                if current_chunk:
                    chunk_text = ' '.join(current_chunk)
                    quality_score = self._calculate_chunk_quality(chunk_text)
                    chunks.append({
                        'text': chunk_text,
                        'quality_score': quality_score,
                        'word_count': len(chunk_text.split()),
                        'sentences': len(current_chunk)
                    })
            else:
                quality_score = self._calculate_chunk_quality(para)
                chunks.append({
                    'text': para,
                    'quality_score': quality_score,
                    'word_count': len(para.split()),
                    'sentences': para.count('.') + para.count('!') + para.count('?')
                })
        # Filter by quality and minimum length
        return [chunk for chunk in chunks 
                if chunk['word_count'] >= MIN_CHUNK_LENGTH 
                and chunk['quality_score'] >= 0.3]

    def _calculate_chunk_quality(self, chunk: str) -> float:
        """Calculate quality score for a chunk"""
        score = 0.0
        
        # Length score (prefer medium-length chunks)
        word_count = len(chunk.split())
        if 20 <= word_count <= 200:
            score += 0.3
        elif word_count > 200:
            score += 0.1
        
        # Content score (prefer chunks with insurance terms)
        insurance_keywords = ['coverage', 'policy', 'claim', 'benefit', 'exclusion', 'limit', 'premium', 'deductible']
        keyword_count = sum(1 for keyword in insurance_keywords if keyword.lower() in chunk.lower())
        score += min(0.4, keyword_count * 0.1)
        
        # Structure score (prefer complete sentences)
        sentence_count = len(re.split(r'[.!?]', chunk))
        if sentence_count >= 2:
            score += 0.2
        
        # Readability score (avoid chunks with too many numbers/special chars)
        special_char_ratio = len(re.findall(r'[^\w\s]', chunk)) / len(chunk)
        if special_char_ratio < 0.3:
            score += 0.1
        
        return min(1.0, score)

    def store_chunks(self, chunks: List[Dict], metadatas: List[Dict]):
        """Enhanced chunk storage with quality filtering"""
        # Filter chunks by quality
        high_quality_chunks = [chunk for chunk in chunks if chunk['quality_score'] >= 0.5]
        
        if not high_quality_chunks:
            # Fallback to all chunks if none meet quality threshold
            high_quality_chunks = chunks
        
        # Prepare for storage
        texts = [chunk['text'] for chunk in high_quality_chunks]
        embeddings = self.embedding_model.encode(texts).tolist()
        
        # Enhanced metadata
        enhanced_metadatas = []
        for i, chunk in enumerate(high_quality_chunks):
            metadata = metadatas[i] if i < len(metadatas) else {}
            enhanced_metadata = {
                **metadata,
                'quality_score': chunk['quality_score'],
                'word_count': chunk['word_count'],
                'sentence_count': chunk['sentences'],
                'processed_at': datetime.now().isoformat()
            }
            enhanced_metadatas.append(enhanced_metadata)
        
        self.collection.add(
            documents=texts,
            embeddings=embeddings,
            metadatas=enhanced_metadatas,  # type: ignore
            ids=[str(uuid.uuid4()) for _ in texts]
        )

    def process_pdf(self, pdf_path: Path):
        """Enhanced PDF processing with LLM metadata extraction"""
        print(f"Processing PDF: {pdf_path.name}")
        
        # Extract text
        with pdfplumber.open(pdf_path) as pdf:
            full_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        
        # Fallback to OCR if text extraction is poor
        if len(full_text.strip()) < 100:
            print(f"Using OCR for {pdf_path.name}")
            full_text = self.ocr_pdf(pdf_path)
        
        cleaned = self.clean_text(full_text)
        if not cleaned.strip():
            print(f"No text extracted from {pdf_path.name}")
            return None
        
        # Enhanced chunking
        chunks = self.semantic_chunking(cleaned)
        if not chunks:
            print(f"No quality chunks from {pdf_path.name}")
            return None
        
        # Enhanced metadata extraction
        section_headers = self.extract_insurance_section_headers(cleaned)
        effective_dates = self.extract_effective_dates(cleaned)
        footnotes = self.extract_footnotes(cleaned)
        tables = self.extract_tables(pdf_path)
        
        # LLM-powered entity extraction
        insurance_entities = self.extract_insurance_entities(cleaned[:2000])  # Limit for LLM
        
        def safe_join(val):
            if isinstance(val, list):
                return "; ".join(str(v) for v in val) if val else ""
            return str(val) if val is not None else ""
        
        # Enhanced metadata
        metadatas = [{
            "source": str(pdf_path.name),
            "page": "aggregated",
            "chunk_id": str(uuid.uuid4())[:8],
            "section_headers": safe_join(section_headers),
            "effective_dates": safe_join(effective_dates),
            "footnotes": safe_join(footnotes),
            "tables": safe_join(tables),
            "policy_types": safe_join(insurance_entities.get("policy_types", [])),
            "coverage_areas": safe_join(insurance_entities.get("coverage_areas", [])),
            "exclusions": safe_join(insurance_entities.get("exclusions", [])),
            "limits": safe_join(insurance_entities.get("limits", [])),
            "waiting_periods": safe_join(insurance_entities.get("waiting_periods", [])),
            "processing_method": "enhanced_pdf"
        } for _ in chunks]
        
        self.store_chunks(chunks, metadatas)
        print(f"Processed {len(chunks)} chunks from {pdf_path.name}")
        return len(chunks)

    def process_docx(self, docx_path: Path):
        """Enhanced DOCX processing"""
        print(f"Processing DOCX: {docx_path.name}")
        
        doc = Document(str(docx_path))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        cleaned = self.clean_text(full_text)
        
        if not cleaned.strip():
            print(f"No text extracted from {docx_path.name}")
            return None
        
        chunks = self.semantic_chunking(cleaned)
        if not chunks:
            print(f"No quality chunks from {docx_path.name}")
            return None
        
        section_headers = self.extract_insurance_section_headers(cleaned)
        effective_dates = self.extract_effective_dates(cleaned)
        footnotes = self.extract_footnotes(cleaned)
        tables = self.extract_tables_docx(doc)
        insurance_entities = self.extract_insurance_entities(cleaned[:2000])
        
        def safe_join(val):
            if isinstance(val, list):
                return "; ".join(str(v) for v in val) if val else ""
            return str(val) if val is not None else ""
        
        metadatas = [{
            "source": str(docx_path.name),
            "page": "aggregated",
            "chunk_id": str(uuid.uuid4())[:8],
            "section_headers": safe_join(section_headers),
            "effective_dates": safe_join(effective_dates),
            "footnotes": safe_join(footnotes),
            "tables": safe_join(tables),
            "policy_types": safe_join(insurance_entities.get("policy_types", [])),
            "coverage_areas": safe_join(insurance_entities.get("coverage_areas", [])),
            "exclusions": safe_join(insurance_entities.get("exclusions", [])),
            "limits": safe_join(insurance_entities.get("limits", [])),
            "waiting_periods": safe_join(insurance_entities.get("waiting_periods", [])),
            "processing_method": "enhanced_docx"
        } for _ in chunks]
        
        self.store_chunks(chunks, metadatas)
        print(f"Processed {len(chunks)} chunks from {docx_path.name}")
        return len(chunks)

    def process_email(self, email_path: Path):
        """Enhanced email processing"""
        print(f"Processing email: {email_path.name}")
        
        text = ""
        tables = []
        
        if email_path.suffix.lower() == ".msg":
            msg = extract_msg.Message(str(email_path))
            text = msg.body or ""
        elif email_path.suffix.lower() == ".eml":
            with open(email_path, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)  # type: ignore
                try:
                    text = msg.get_body(preferencelist=('plain', 'html')).get_content()  # type: ignore
                except Exception:
                    if msg.is_multipart():
                        text = "\n".join(part.get_payload(decode=True).decode(errors='ignore')  # type: ignore
                                             for part in msg.walk() if part.get_content_type() == 'text/plain')
                    else:
                        text = msg.get_payload(decode=True).decode(errors='ignore') if msg.get_payload() else ""  # type: ignore
        
        cleaned = self.clean_text(text)
        if not cleaned.strip():
            print(f"No text extracted from {email_path.name}")
            return None
        
        chunks = self.semantic_chunking(cleaned)
        if not chunks:
            print(f"No quality chunks from {email_path.name}")
            return None
        
        section_headers = self.extract_insurance_section_headers(cleaned)
        effective_dates = self.extract_effective_dates(cleaned)
        footnotes = self.extract_footnotes(cleaned)
        insurance_entities = self.extract_insurance_entities(cleaned[:2000])
        
        def safe_join(val):
            if isinstance(val, list):
                return "; ".join(str(v) for v in val) if val else ""
            return str(val) if val is not None else ""
        
        metadatas = [{
            "source": str(email_path.name),
            "page": "aggregated",
            "chunk_id": str(uuid.uuid4())[:8],
            "section_headers": safe_join(section_headers),
            "effective_dates": safe_join(effective_dates),
            "footnotes": safe_join(footnotes),
            "tables": safe_join(tables),
            "policy_types": safe_join(insurance_entities.get("policy_types", [])),
            "coverage_areas": safe_join(insurance_entities.get("coverage_areas", [])),
            "exclusions": safe_join(insurance_entities.get("exclusions", [])),
            "limits": safe_join(insurance_entities.get("limits", [])),
            "waiting_periods": safe_join(insurance_entities.get("waiting_periods", [])),
            "processing_method": "enhanced_email"
        } for _ in chunks]
        
        self.store_chunks(chunks, metadatas)
        print(f"Processed {len(chunks)} chunks from {email_path.name}")
        return len(chunks)

def main():
    processor = EnhancedDocumentProcessor()
    data_dir = Path(__file__).parent.parent / "data"
    
    total_chunks = 0
    for file_path in data_dir.glob("*.*"):
        try:
            if file_path.suffix.lower() == ".pdf":
                print(f"Processing {file_path.name} (PDF)...")
                chunk_count = processor.process_pdf(file_path)
                if chunk_count:
                    total_chunks += chunk_count
            elif file_path.suffix.lower() == ".docx":
                print(f"Processing {file_path.name} (Word)...")
                chunk_count = processor.process_docx(file_path)
                if chunk_count:
                    total_chunks += chunk_count
            elif file_path.suffix.lower() in [".eml", ".msg"]:
                print(f"Processing {file_path.name} (Email)...")
                chunk_count = processor.process_email(file_path)
                if chunk_count:
                    total_chunks += chunk_count
        except Exception as e:
            print(f"Error processing {file_path.name}: {str(e)}")
    
    print(f"Total chunks processed: {total_chunks}")

if __name__ == "__main__":
    main()
